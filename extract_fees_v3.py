#!/usr/bin/env python3
"""Extract pre-KG fee structure from school PDFs - Version 3 (Refined)"""

import re
from pathlib import Path
from datetime import datetime
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_all_text_from_pdf(pdf_path):
    """Extract all text from all pages of PDF"""
    try:
        reader = PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        print(f"  ✗ Error reading PDF: {e}")
        return ""

def find_pre_kg_fee_info(pdf_path, school_name):
    """Extract pre-KG fee information from PDF with improved logic"""
    text = extract_all_text_from_pdf(pdf_path)
    text_lower = text.lower()

    info = {
        "school_name": school_name,
        "level": None,
        "fees": None,
        "duration_type": None,
        "confidence": "Low",
        "comments": ""
    }

    # Priority-based level identification
    pre_kg_patterns = [
        (r"pre\s*-?\s*k[g1]|pre\s*-?\s*kindergarten", "Pre-KG"),
        (r"early\s+years|early\s+start|toddler", "Early Years"),
        (r"reception", "Reception"),
        (r"nursery|p1\s*-\s*p3", "Nursery"),
    ]

    found_level = False
    for pattern, level_name in pre_kg_patterns:
        if re.search(pattern, text_lower):
            info["level"] = level_name
            info["confidence"] = "High" if "pre-kg" in pattern else "Medium"
            found_level = True
            break

    if not found_level:
        info["comments"] = "Pre-KG level not found. Manual review required."
        return info

    # Extract fees - look for common fee keywords
    fee_keywords = [
        r"tuition\s+fee[:\s]+[\d,]+",
        r"annual\s+(?:education\s+)?fee[:\s]+[\d,]+",
        r"education\s+fee[:\s]+[\d,]+",
        r"total\s+fee[:\s]+[\d,]+",
    ]

    fees_found = []
    for keyword_pattern in fee_keywords:
        matches = re.findall(keyword_pattern, text_lower)
        if matches:
            # Extract the actual numbers
            for match in matches:
                numbers = re.findall(r'[\d,]+', match)
                if numbers:
                    fees_found.append({
                        'keyword': match.split(':')[0].strip(),
                        'amount': numbers[0]
                    })

    # Filter for reasonable fees (1000-1000000 INR range for pre-KG)
    reasonable_fees = []
    for fee_info in fees_found:
        try:
            fee_num = int(fee_info['amount'].replace(',', ''))
            if 1000 < fee_num < 1000000:
                reasonable_fees.append(fee_info)
        except:
            pass

    if reasonable_fees:
        # Prefer "Tuition Fee" or "Annual Fee"
        preferred_fee = None
        for fee_info in reasonable_fees:
            if 'tuition' in fee_info['keyword']:
                preferred_fee = fee_info['amount']
                break
        if not preferred_fee:
            preferred_fee = reasonable_fees[0]['amount']

        info["fees"] = preferred_fee
        info["confidence"] = "High"
    else:
        # Try extracting any reasonable number near pre-kg keyword
        level_lower = info["level"].lower()
        idx = text_lower.find(level_lower)
        if idx != -1:
            context = text[max(0, idx-300):min(len(text), idx+1000)]
            amounts = re.findall(r'[\d,]+(?:\.\d{2})?', context)
            for amount in amounts:
                try:
                    num = int(amount.replace(',', '').split('.')[0])
                    if 10000 < num < 1000000:
                        info["fees"] = amount
                        info["confidence"] = "Medium"
                        info["comments"] = "Fee extracted from context; manual verification recommended."
                        break
                except:
                    pass

    # Detect duration type
    if re.search(r'\bper\s+annum\b|\bannual\b|per\s+academic\s+year', text_lower):
        info["duration_type"] = "Annual"
    elif re.search(r'\bper\s+term\b|per\s+semester|term\s+fee', text_lower):
        info["duration_type"] = "Per term"

    # Check for USD currency
    if '$' in text and 'usd' in text_lower:
        if info["fees"]:
            info["comments"] = "Fees in USD. Manual conversion required."

    return info

def create_docx(output_path, data_list):
    """Create the information.docx file"""
    doc = Document()

    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run("School Pre-KG Fee Structure - Bangalore")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add timestamp
    timestamp = doc.add_paragraph(f"Last Updated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    timestamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add data description
    desc = doc.add_paragraph("Extracted from school fee structure PDFs. Manual verification recommended for accuracy.")
    desc.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Create table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'

    # Set column widths
    for row in table.rows:
        row.cells[0].width = 2500000
        row.cells[1].width = 1500000
        row.cells[2].width = 1500000
        row.cells[3].width = 1500000
        row.cells[4].width = 1200000
        row.cells[5].width = 2000000

    # Add header row
    header_cells = table.rows[0].cells
    headers = ['School Name', 'Level', 'Fees (INR)', 'Duration Type', 'Confidence', 'Comments']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add data rows
    for item in data_list:
        row_cells = table.add_row().cells
        row_cells[0].text = item.get('school_name', '')
        row_cells[1].text = item.get('level', '') or 'NOT FOUND'
        row_cells[2].text = item.get('fees', '') or 'NOT FOUND'
        row_cells[3].text = item.get('duration_type', '') or 'NOT SPECIFIED'
        row_cells[4].text = item.get('confidence', 'Low')
        row_cells[5].text = item.get('comments', '')

    doc.save(output_path)
    print(f"\n✓ Document created: {output_path}")

def main():
    # Paths - works when run from the script directory
    script_dir = Path(__file__).parent
    brochures_dir = script_dir
    output_file = script_dir / "information.docx"

    print("=" * 80)
    print("SCHOOL FEE EXTRACTION")
    print("=" * 80)
    print(f"Scanning: {brochures_dir}\n")

    # Find all PDFs sorted by date modified
    pdf_files = sorted(brochures_dir.glob("*.pdf"), key=lambda x: x.stat().st_mtime, reverse=True)
    print(f"Found {len(pdf_files)} PDF files\n")

    # School name mapping
    school_mapping = {
        "BIS": "Bangalore International School",
        "IISB": "Indus International School Bangalore",
        "Indian_": "Stonehill International School",
        "Neev": "Neev Academy"
    }

    data_list = []

    for pdf_path in pdf_files:
        print(f"Processing: {pdf_path.name}")

        # Determine school name from filename
        school_name = pdf_path.stem
        for key, val in school_mapping.items():
            if key in school_name:
                school_name = val
                break

        # Extract info
        info = find_pre_kg_fee_info(str(pdf_path), school_name)
        print(f"  ✓ School: {info['school_name']}")
        print(f"  ✓ Level: {info['level'] or 'NOT FOUND'}")
        print(f"  ✓ Fees: {info['fees'] or 'NOT FOUND'}")
        print(f"  ✓ Duration: {info['duration_type'] or 'NOT SPECIFIED'}")
        print(f"  ✓ Confidence: {info['confidence']}")
        if info['comments']:
            print(f"  ⚠ Comments: {info['comments']}")
        print()

        data_list.append(info)

    # Create document
    print("Creating DOCX file...")
    create_docx(str(output_file), data_list)
    print("=" * 80)
    print("✓ COMPLETE!")
    print("=" * 80)

if __name__ == "__main__":
    main()
